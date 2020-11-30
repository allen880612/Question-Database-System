import docx

print("Start")
word = docx.Document()
word.add_heading("Test", 0) #新增那個醜醜藍字

#新增題目 style
questionStyle = word.styles.add_style("question", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH) #新增一樣式 (樣式名稱, 樣式類型)
questionStyle.font.size = docx.shared.Pt(12) #更改此樣式的文字大小
questionStyle.font.name = "Times New Roman" #設定英文字體
# where am I? who am I???
questionStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #設定中文字體
# 設定凸排, su go i ne, my Python
questionStyle.paragraph_format.first_line_indent = docx.shared.Pt(-18) # 設定首縮排/凸排 (正值 = 縮排, 負值 = 凸排)
questionStyle.paragraph_format.left_indent = docx.shared.Pt(18) # ↓注意，重點來了，設定"整個段落"縮排  (正常來說應該不用設定，但是設定凸排的時候，他會順便把整個段落也往左移動，所以要他媽的移回來)


#新增題目
for i in range(0, 10):
    questionIndex = "(" + str(i + 1) + ") " #題號
    paragraph = word.add_paragraph(questionIndex + str(i), style = "question") #題號 + 題目 一題作為一個段落
    paragraph.alignment = 3 #設定段落對齊 0 = 靠左, 1 = 置中, 2 = 靠右, 3 = 左右對齊 (WD_PARAGRAPH_ALIGNMENT)

#word.add_page_break() #應該是強制換頁

run = word.paragraphs[5].add_run()
#run.add_picture("miku.jpg", height=docx.shared.Cm(4), width=docx.shared.Cm(4))
run.add_picture("database\\數學\\應用題\\典型應用題\\燕尾題\\1.png")
#word.add_picture("miku.jpg", height=docx.shared.Cm(4), width=docx.shared.Cm(4))


word.save("picTest.docx") #存檔 (存在word資料夾)

#print( word.paragraphs[5].text)
print("Done")
