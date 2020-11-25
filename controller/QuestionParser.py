from docx import Document
from lxml import etree
import asciitomathml.asciitomathml

'''
使用流程
1. 建立 QuestionParserf 的Instance，必須要呼叫 Initialize後才可使用
2. Initialize 傳入 [原始題目字串，該題目段落，(包夾特殊算是的字元)]
3. 呼叫 GetMathEquationParagraph()，取得處理好的 paragraph

註：建立一次後，2, 3可重複呼叫使用，但注意3 一定要在呼叫 2後使用
'''

class QuestionParser():
	def __init__(self):
		pass
		# 注意，必須要初始化才能使用

	# 初始化
	def Initialize(self, src_string=str, paragraph=None, hint_str='♥'):
		self.Src_str = src_string
		self.Paragraph = paragraph
		self.Hint_str = hint_str

	def Reset(self):
		self.Src_str = ''
		self.Paragraph = None
		self.Hint_str = '♥'

	def GetMathEquationParagraph(self):
		self.ParseQuestion()
		return self.Paragraph


	def ParseQuestion(self):
		isSymbolPart = False
		temp_symbol_string = ''

		for ch in self.Src_str:
			if isSymbolPart:
				if ch == self.Hint_str:
					# 結束算式(特殊符號)區域 > encode > append to paragraph
					self.__ProcessSymbolPart(temp_symbol_string)
					temp_symbol_string = ''
					isSymbolPart = False
				else:
					# 符號段落中，紀錄至temp
					temp_symbol_string += ch
			else:
				if ch == self.Hint_str:
					# 開始記錄符號
					isSymbolPart = True
				else:
					# 一般文字段落
					self.Paragraph.add_run(ch)

	def __ProcessSymbolPart(self, symbol_str=str):

		mathEquation = self.GetMathEquationByString( symbol_str )
		self.Paragraph._element.append( mathEquation )

	
	def GetMathEquationByString( self, equation_string=str ):
		#question_string = '1. x=(-b±sqrt(b^2-4ac))/(2a), 解b'
	
		# -預處理 
		equation_string = self.__PreProcess( equation_string )
		
		# -儲存線段資訊
		line_list = self.__GetLineList( equation_string )

		# Unicode -> xml_string
		math_obj =  asciitomathml.asciitomathml.AsciiMathML()
		math_obj.parse_string( equation_string )
		xml_string = math_obj.to_xml_string(encoding='unicode')

		# -後處理替換線段 (之後向量 的處理方式應該也可以照抄)
		for line in line_list:
			xml_string = xml_string.replace( self.__GetCompleteOverLineKeyWord( line ), self.__GetXMLOverline( line ) )

		return self.__GetMathEquationByXMLString( xml_string )
	
	'''
	透過 p = document.add_paragraph('')
	p._element.append( 回傳值 ) 來使用
	'''
	@classmethod
	def __GetMathEquationByXMLString( clc, xml_string ):
		tree = etree.fromstring(xml_string)
		xslt = etree.parse('MML2OMML.XSL')	# 注意這個檔案是不是在預期路徑上
		transform = etree.XSLT(xslt)
		new_dom = transform(tree)
		return new_dom.getroot()

	@classmethod
	def __PreProcess(clc, equation_string=str):
		def PreProcessLine( question_str=str ):
			new_question = question_str.replace( '("', '' )
			new_question = new_question.replace( '" )', '' )
			return new_question

		equation_string = PreProcessLine( equation_string )
		equation_string = equation_string.replace('√', 'sqrt')

		return equation_string

	@classmethod
	def __GetLineList( clc, question_string=str ):
		line_list = list()
		temp_line = ''
		overline_flag = False;
		index = 0

		for s in question_string:
			index += 1
			if s == '¯':
				#temp_line = s
				overline_flag = True

			elif overline_flag:
				if s.encode('utf-8').isalpha():
					temp_line += s
				else:
					overline_flag = False
					line_list.append(temp_line)
					temp_line = ''

		# (若有的話)補上結尾的線段
		if temp_line:
			line_list.append(temp_line)

		print (line_list)
		return line_list

	'''
	處理線段，把 (¯BD) 包裝成 :
	<mo>¯</mo><mi>B</mi><mi>D</mi>
	'''
	@classmethod
	def __GetCompleteOverLineKeyWord(clc, line_str=str ):
		returner = '<mo>¯</mo>'

		for s in line_str:
			returner += '<mi>' + s + '</mi>'

		return returner

	'''
	針對線段的處理function
	線段BC
	複製到的會是 : ¯BC
	mathML 表達為 : <mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mover accent="false"><mml:mrow><mml:mi>B</mml:mi><mml:mi>D</mml:mi></mml:mrow><mml:mo>¯</mml:mo></mml:mover></mml:math>
	'''
	@classmethod
	def __GetXMLOverline( clc, line_str=str ):
		BEGIN = '<mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mover accent="false"><mml:mrow>'
		END = '</mml:mrow><mml:mo>¯</mml:mo></mml:mover></mml:math>'
		line = ''

		for s in line_str:
			line += '<mml:mi>' + s + '</mml:mi>'
		return BEGIN + line + END

'''
Test Case
'''
if __name__ == '__main__':
	
	d = Document()
	p = d.add_paragraph("")

	q1 = '1.♥x=(-b±√(b^2-4ac))/(2a)♥, 請問 a, b = ?'
	q2 = '2.若♥¯BD：¯CD♥＝2：5，♥¯("AE" )：¯("DE" )♥＝1：4，則△ABE面積：△CDE面積＝【 1：10 】'

	qp = QuestionParser(q1, p)
	qp.ParseQuestion()

	p2 = d.add_paragraph("")
	qp.Initialize(q2, p2)
	qp.ParseQuestion()

	d.save("simpleEq.docx")