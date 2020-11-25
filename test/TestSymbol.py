from docx import Document
from lxml import etree
import asciitomathml.asciitomathml

'''
針對數學算式做處理，傳入要加入的段落，題目string (也可只傳運算式的部分，看需求)
但傳入的都會透用預設的斜體，以及運算式的字形，這部分需要額外處理；
或是就只針對算式的部分再做轉換

已知限制如下:
1. 分子、母需手動補上沒有的括號，如: a/2b 必須 至少要是 a/(2b)，否則會分析成 a/2 * b
2. 跟號需要額外處理 √(b^2-4ac) 必須 轉成 sqrt(b^2-4ac)
3. 線段、向量 需要透過HardCode的方式幫他補齊轉換，請見 ConvertOverline2MathML
'''
def GetMathEquationByString( question_string=str ):
	#question_string = '1. x=(-b±sqrt(b^2-4ac))/(2a), 解b'
	
	# -預處理 
	question_string = PreProcessLine( question_string )
	question_string = question_string.replace('√', 'sqrt')
	# -儲存線段資訊
	line_list = GetLineList( question_string )

	# Unicode -> xml_string
	math_obj =  asciitomathml.asciitomathml.AsciiMathML()
	math_obj.parse_string( question_string )
	xml_string = math_obj.to_xml_string(encoding='unicode')

	# -後處理替換線段 (之後向量 的處理方式應該也可以照抄)
	for line in line_list:
		xml_string = xml_string.replace( GetCompleteOverLineKeyWord( line ), GetXMLOverline( line ) )

	return GetMathEquationByXMLString( xml_string )


# 針對線段的前處理
def PreProcessLine( question_str=str ):
	new_question = question_str.replace( '("', '' )
	new_question = new_question.replace( '" )', '' )
	return new_question

def GetLineList( question_string=str ):
	line_list = list()
	temp_line = ''
	overline_flag = False;

	for s in question_string:
		if s == '¯':
			#temp_line = s
			overline_flag = True
		elif overline_flag:
			if s.encode('utf-8').isalpha():
				temp_line += s
			else:
				overline_flag = False
				line_list.append(temp_line)
				temp_line = ''

	print (line_list)
	return line_list

'''
針對線段的處理function
線段BC
複製到的會是 : ¯BC
mathML 表達為 : <mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mover accent="false"><mml:mrow><mml:mi>B</mml:mi><mml:mi>D</mml:mi></mml:mrow><mml:mo>¯</mml:mo></mml:mover></mml:math>
'''
def GetXMLOverline( line_str=str ):
	BEGIN = '<mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mover accent="false"><mml:mrow>'
	END = '</mml:mrow><mml:mo>¯</mml:mo></mml:mover></mml:math>'
	line = ''

	for s in line_str:
		line += '<mml:mi>' + s + '</mml:mi>'
	return BEGIN + line + END

'''
♥ 愛戰勝一切，處理算式部分 ♥
'''
def BulidFormulaDict( question_string=str, hint_str='♥' ):
	formula_dict = {}
	temp_formula = ''
	flag = False;

	for i in range(0, len(question_string)):
		s = question_string[i]
		if s == hint_str:
			flag = not flag
		elif flag:
			temp_formula += s
			if i+1 < len(question_string) and question_string[i+1] == hint_str:
				formula_dict[temp_formula] = ''
				temp_formula = ''
			
	for formula_str in formula_dict:
		formula_dict[ formula_str ] = GetMathEquationByString( formula_str )
	#print (formula_dict)
	return formula_dict

'''
 formula字典 原字串-> 對應 XML encode 函式
 線段list
'''
def EncodeSourceQuestion( question_string=str, hint_str='♥' ):
	# 預處理 
	question_string = PreProcessLine( question_string )
	question_string = question_string.replace('√', 'sqrt')
	# 儲存線段資訊
	line_list = GetLineList( question_string )
	#儲存函式字典
	formula_dict = BulidFormulaDict( question_string )


'''
處理線段，把 (¯BD) 包裝成 :
<mo>¯</mo><mi>B</mi><mi>D</mi>
'''
def GetCompleteOverLineKeyWord( line_str=str ):
	returner = '<mo>¯</mo>'

	for s in line_str:
		returner += '<mi>' + s + '</mi>'

	return returner


'''
# etree 神秘轉換成word格式

傳入xml 形式的算式，大概長這樣:

<math xmlns="http://www.w3.org/1998/Math/MathML">
<mstyle><mi>x</mi><mo>=</mo><mfrac><mrow class="nominator"><mo>-</mo><mi>b</mi><mo>±</mo>
<msqrt><mrow class="radical"><msup><mi>b</mi><mn>2</mn></msup><mo>-</mo><mn>4</mn><mi>a</mi><mi>c</mi></mrow></msqrt>
</mrow><mrow class="denominator"><mn>2</mn><mi>a</mi></mrow></mfrac></mo></mstyle></math>)

透過 p = document.add_paragraph('')
p._element.append( 回傳值 ) 來使用
'''
def GetMathEquationByXMLString( xml_string ):
	tree = etree.fromstring(xml_string)
	xslt = etree.parse('MML2OMML.XSL')	# 注意這個檔案是不是在預期路徑上
	transform = etree.XSLT(xslt)
	new_dom = transform(tree)
	return new_dom.getroot()


# 建 & 寫 Word
document = Document()
Titleparagraph = document.add_paragraph("測試數學等各種運算")
# 測試線段
testLineP = document.add_paragraph("測試線段:")
overline_AB_xml = GetXMLOverline('AB')
testLineP._element.append( GetMathEquationByXMLString( overline_AB_xml ) )
testLineQuestionP = document.add_paragraph("測試線段題目:")
overline_question_string = '(1)若¯BD：¯CD＝2：5，¯("AE" )：¯("DE" )＝1：4，則△ABE面積：△CDE面積＝【 1：10 】'
#overline_question_string = '(1)若♥√(b^2-4ac)♥¯BD：¯CD＝2：5，¯("AE" )：¯("DE" )＝1：4，♥x=(-b±√(b^2-4ac))/(2a)♥則△ABE面積：△CDE面積＝【 1：10 】'
testLineQuestionP._element.append( GetMathEquationByString( overline_question_string ) )

# 測試 跟號, 分數, 次方
testSqrtFracPowP = document.add_paragraph("測試跟號, 分數, 次方:")
question_string = '1.♥x=(-b±√(b^2-4ac))/(2a)♥, 請問 a, b = ?'
testSqrtFracPowP._element.append( GetMathEquationByString( question_string ) )
#p.add_run(' = 題目後續')
document.save("simpleEq.docx")

'''3¯BD＝2¯CD
latex & sympy 作法，大同小異，
但先用sympy 表捯後 , 再pasrse 成latex 成本、難度高，也沒必要，
但不排除在Word端用MathType 或類似套件先轉譯好，再用 latex_to_text 轉 XML
'''
from sympy import *
from pylatexenc.latexencode import unicode_to_latex
from pylatexenc.latex2text import LatexNodes2Text

#latexEq = unicode_to_latex(u'(a+1)/b = 2a')
#print(latexEq)

'''
最大的問題在於無法處理:
latex = latex(sympify('x+(-b±√(b^2-4ac))/2a'))
像 4ac 這種形式的參數
故若無法在Word端就先轉換為latex形式，就不考慮這種處理方式
'''

latex = latex(sympify('(sqrt(a)+1)/b'))
#latex = r"\frac {1 + \sqrt {\a}} {\b}"
text = LatexNodes2Text().latex_to_text(latex)
#pprint(latex, use_unicode=True)
#print (text)

# create expression by sympy
x, y = symbols('x y')
#expr1 = (x+y)**2
# expr1 = x/y

from sympy.parsing.latex import parse_latex
expr1 = parse_latex(latex)
#expr1 = parse_latex(r"\frac {1 + \sqrt {\a}} {\b}")

# create MathML structure
expr1xml = mathml(expr1, printer = 'presentation')
#P._element.append( GetMathEquationByXMLString( expr1xml ) )

