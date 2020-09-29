# -*- coding: utf-8 -*-
from docx import Document
import sys

document = Document()

# 傳入所在段落，與要上標處理的字串
def AddUpperWord( paragraph , upperText):
	upper_text = paragraph.add_run( str(upperText) )
	upper_text.font.superscript = True

def AddOverLine( paragraph, text=str ):
	retruner = '' 
	#retruner = u'%s' % text
	#retruner += u' \u0305'*len(text)
	for t in str(text):
		#retruner += t + '\u0305'
		#retruner += u'%s' % t + '\u0304'
		retruner += t + u'\u0305'
	paragraph.add_run(retruner)
	print(retruner)

'''
	Test function
'''
def TestUpper():
	p = document.add_paragraph('X')
	AddUpperWord(p, 2)

	p.add_run(' + Y')
	AddUpperWord(p, 2)
	p.add_run(' = 214')

def TestOverlines():
	p = document.add_paragraph('若')
	AddOverLine( p, u'CD' )
	p.add_run(' = 214')


TestOverlines()

try:
	document.save('TestOverline.docx')
	print ('done')
except:
	print (sys.exc_info()[0])
