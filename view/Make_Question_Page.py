from PyQt5.QtWidgets import QMainWindow, QMessageBox
from PyQt5.QtGui import QFont
from PyQt5 import QtCore, QtWidgets, QtGui
from view.UI import Make_Question_UI_Test as mq_UI_Test
from view.UI import Make_Question_UI as mq_UI
from view import Revise_MakeQuestion_Page
from model import MyLibrary
import random
import os
import docx
from functools import partial

class MakeQuestionPage(QMainWindow):

    # 返回首頁Signal
    return_mainpage_signal = QtCore.pyqtSignal(bool) # set 信號

    #region 變數宣告
    TEMPLATE_WORD_PATH = "database/default.docx"
    comboboxSelectOption = []
    #endregion

    def __init__(self, _model):
        super(MakeQuestionPage, self).__init__()

        self.ui = mq_UI.MakeQuestionPage_UI()
        self.ui.setupUi(self)

        self.model = _model

        # 上一層所選的level的list
        self.question_level_list = []

        # UI上擺放題目的Layout
        self.question_layout = self.ui.gridLayout_1

        # UI上 的題目文字
        self.question_label = []

        # UI上 的題數輸入框
        self.question_tbox = []

        # dict 用來拿最大題數
        self.question_dict = {}

        # 微調題目頁面
        self.Revise_MakeQuestion_View = Revise_MakeQuestion_Page.ReviseMakeQuestionPage(self.model)
        self.Is_revise_question_view_open = False

        self.Initialize()

    # 初始化
    def Initialize(self):
        self.ResetPage()
        self.ConnectEvent()

    # 註冊事件
    def ConnectEvent(self):
        #self.ui.button_make_question.clicked.connect(self.Confirm)
        self.ui.button_make_question.clicked.connect(self.OpenReviseQuestionView)
        self.ui.lineEdit_total_number.setValidator(QtGui.QIntValidator(self.ui.lineEdit_total_number)) # 設置總題數只能輸入數字
        self.ui.lineEdit_total_number.editingFinished.connect(self.InputOnTotleNumber) # LineEdit事件 - 輸入完後均等題目數
        
        # 註冊 調整題目頁面之信號
        self.Revise_MakeQuestion_View.revise_make_question_signal.connect(self.GetReviseMakeQuestionViewData)

    # 重設頁面
    def ResetPage(self):
        # 重設defaultString
        defaultString = self.model.DefaultString_NoSelect
        
        # 重建題目數
        self.DeletaAllQuestion()
        for question in self.question_level_list:
            self.SetQuestion(question)

        self.question_dict.clear()
        for level in self.question_level_list:
            d_key = MyLibrary.CreateDictKey(level)
            self.question_dict[d_key] = self.model.GetQuestionList(level)

        # 重置總題數 & 等分題數
        self.ui.lineEdit_total_number.setText("40")
        self.InputOnTotleNumber() # -> 移到show Event
        self.Is_revise_question_view_open = False

    # 刪除所有題目
    def DeletaAllQuestion(self):
        for label in self.question_label:
            self.question_layout.removeWidget(label)
            label.setVisible(False)
        for tbox in self.question_tbox:
            self.question_layout.removeWidget(tbox)
            tbox.setVisible(False)
        self.question_label.clear()
        self.question_tbox.clear()

    # 設置一道題目
    def SetQuestion(self, add_question_strlist, question_number=0):
        # self.gridLayout_1.addWidget(self.label_total_number, 0, 2, 1, 1)
        # self.gridLayout_1.addWidget(self.lineEdit_total_number, 0, 3, 1, 1)
        current_question_number = len(self.question_label)
        add_row = current_question_number + 1 # 要新增到的列的位置

        # Set Font
        font = QFont()
        font.setFamily("微軟正黑體")
        font.setPointSize(15)

        # Set Label
        new_label = QtWidgets.QLabel(self.ui.centralwidget)
        new_label.setFont(font)
        new_label.setText(MyLibrary.GetQuestionShowText(add_question_strlist))
        new_label.setObjectName(self.GetNewObjectName("label"))
        self.question_layout.addWidget(new_label, add_row, 2, 1, 1)

        # Set Tbox
        new_tbox = QtWidgets.QLineEdit(self.ui.centralwidget)
        new_tbox.setFont(font)
        new_tbox.setText(str(question_number))
        new_tbox.setObjectName(self.GetNewObjectName("textbox"))
        self.question_layout.addWidget(new_tbox, add_row, 3, 1, 1)
        new_tbox.setValidator(QtGui.QIntValidator(new_tbox)) # 設置只能輸入數字
        new_tbox.editingFinished.connect(self.InputQuestionNumberEvent)
        
        self.question_label.append(new_label)
        self.question_tbox.append(new_tbox)

    # 得到新元件的物件名
    def GetNewObjectName(self, type):
        return type + "_" + str(len(self.question_label))

    # 平均每一題
    def AverageQuestionNumber(self):
        total_question_number = int(self.ui.lineEdit_total_number.text())
        question_level_list_count = len(self.question_level_list) if self.question_level_list else 1 # 上一層選了多少路徑 (至少為1)
        average_number = total_question_number // question_level_list_count # 平均每一個題型有多少題
        remain_number = total_question_number % question_level_list_count  # 多的題目
        average_number_list = [average_number] * question_level_list_count
        for index in range(0, len(self.question_tbox)):
            target_tbox = self.question_tbox[index]
            avg_number = average_number_list[index]
            if remain_number > 0:
                avg_number += 1
                remain_number -= 1
            target_tbox.setText(str(avg_number))
    
    # 將題數轉換成數字list
    def GetQuestionNumberList(self):
        number_list = []
        for tbox in self.question_tbox:
            number_list.append(int(tbox.text()))
        return number_list

    # 輸入題數後發生之事件
    def InputQuestionNumberEvent(self):
        update_total_number = 0
        tbox_count = 0
        for tbox in self.question_tbox:
            qList = self.question_dict[MyLibrary.CreateDictKey(self.question_level_list[tbox_count])]
            max_count = len(qList)
            # 沒有輸入
            if tbox.text() == "":
                tbox.setText("0")
            # 成功輸入
            else:
                # 輸入的超過ㄌ
                if int(tbox.text()) > max_count:
                    tbox.setText(str(max_count))
                    update_total_number += max_count
                    self.ShowTips(MyLibrary.GetQuestionShowText(self.question_level_list[tbox_count]))
                else:
                    update_total_number += int(tbox.text())
            tbox_count += 1
        self.ui.lineEdit_total_number.setText(str(update_total_number))

    # 輸入總題數時發生的事情
    def InputOnTotleNumber(self):
        self.AverageQuestionNumber() # 先均分每一題
        
        # 每一題取最大題數
        update_total_number = 0
        tbox_count = 0
        tips_content = []
        for tbox in self.question_tbox:
            qList = self.question_dict[MyLibrary.CreateDictKey(self.question_level_list[tbox_count])]
            max_count = len(qList)
            if int(tbox.text()) > max_count:
                tbox.setText(str(max_count))
                update_total_number += max_count
                tips_content.append(MyLibrary.GetQuestionShowText(self.question_level_list[tbox_count]))
            else:
                update_total_number += int(tbox.text())
            tbox_count += 1
        if tips_content != []:
            self.ShowTips("\n".join(tips_content))
        self.ui.lineEdit_total_number.setText(str(update_total_number))

    # 接收 來自上一層的題目列表
    def GetQuestionLevelList(self, questionList):
        self.question_level_list = questionList

    # 測試用
    def Test(self):
        qqq = self.model.GetQuestionList(self.question_level_list[0])
        print("-----")
        print(qqq)
        print(len(qqq))

    # Show Tips
    def ShowTips(self, information):
        QMessageBox.information(self, "警告", information + "\n超過最大題數!", QMessageBox.Yes)
    
    ########################################################################
    # 移駕至調整題目頁面
    #region 函式區
    def Confirm(self):
        #excel
        if self.model.IsLoad():
            if not MyLibrary.IskWordOpen("word/answer.docx") or not MyLibrary.IskWordOpen("word/question.docx"):
                QMessageBox.information(self, "警告", "Word開啟中！\n請關閉Word後再試一次", QMessageBox.Yes)
                return

            #questionType = ["數學"]
            qList = [] # 要給Word建構的題目列表
            # 建構 Question List
            for question_level_index in range(0, len(self.question_level_list)):
                this_question_level = self.question_level_list[question_level_index] # 目前跑到的題目階層
                this_question_number = int(self.question_tbox[question_level_index].text()) # 目前跑到的題目階層所要建構的題數
                filter_question_list = self.model.GetQuestionList(this_question_level) # 根據階層篩選題目
                # ↓目前有問題，如果選20題，但該題型只有10題?，目前做法: 直接選10題不重複，但總題數減少
                filter_question_list = random.sample(filter_question_list, min(this_question_number, len(filter_question_list))) # 將題目不重複隨機選擇 k 題 (0 <= k <= 篩選後的題目數量)
                
                qList += filter_question_list

            self.BulidWord(qList, "answer", True)  # 建造word 保留答案
            self.BulidWord(qList, "question", False)  # 建造word 刪除答案
            print("done")
        #excel

    def BulidWord(self, qList, fileName, haveAnswer):
        word = docx.Document(docx=self.TEMPLATE_WORD_PATH)  # 另一個坑，為了讓方裝後也能抓到default.docx，必須指定
        # heading = " - ".join(self.comboboxSelectOption)
        heading = " - ".join(self.question_level_list[0])
        word.add_heading(heading, 0) #新增那個醜醜藍字

        #新增題目 style
        questionStyle = word.styles.add_style("question", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH) #新增一樣式 (樣式名稱, 樣式類型)
        questionStyle.font.size = docx.shared.Pt(12) #更改此樣式的文字大小
        questionStyle.font.name = "Times New Roman" #設定英文字體
        # where am I? who am I???
        questionStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #設定中文字體
        # 設定凸排, su go i ne, my Python
        questionStyle.paragraph_format.first_line_indent = docx.shared.Pt(-18) # 設定首縮排/凸排 (正值 = 縮排, 負值 = 凸排)
        questionStyle.paragraph_format.left_indent = docx.shared.Pt(18) # ↓注意，重點來了，設定"整個段落"縮排  (正常來說應該不用設定，但是設定凸排的時候，他會順便把整個段落也往左移動，所以要他媽的移回來)

        #imageList = self.excel.GetFilteredImage(self.comboboxSelectOption)
        #新增題目
        for i in range(0, len(qList)):
            questionIndex = "(" + str(i + 1) + ") " #題號
            if haveAnswer:
                question = qList[i].GetAnswer()
            else:
                question = qList[i].GetQuestion()
            paragraph = word.add_paragraph(questionIndex + question, style = "question")
            #paragraph = word.add_paragraph(questionIndex + questionList[i], style = "question") #題號 + 題目 一題作為一個段落
            paragraph.alignment = 0 #設定段落對齊 0 = 靠左, 1 = 置中, 2 = 靠右, 3 = 左右對齊 (WD_PARAGRAPH_ALIGNMENT)
            try:
                # print("have", qList[i].HaveImage())
                # print("image", qList[i].GetImage())
                if qList[i].HaveImage() and qList[i].GetImage():
                    print(qList[i].GetImage())
                    #run = word.paragraphs[i + 1].add_run()
                    imageParagraph = word.add_paragraph() # 為了讓圖片靠右，直接新增一個段落，方便用alignment
                    imageParagraph.alignment = 2
                    run = imageParagraph.add_run()
                    #run.add_break() #不換段換行
                    for image in qList[i].GetImage():
                        run.add_picture(image, height=docx.shared.Cm(2.6))
            except:
                print("Insert image fail!")
        
        savePath = "word/" + fileName + ".docx"
        word.save(savePath) #存檔 (存在word資料夾)
    #endregion

    #######################################################

    # 開啟 微調題目 視窗
    def OpenReviseQuestionView(self):
        if self.Is_revise_question_view_open == False:
            self.Is_revise_question_view_open = True
            tuple_list = zip(self.question_level_list, self.GetQuestionNumberList())
            
            self.Revise_MakeQuestion_View.SetQuestionLevelTupleData(list(tuple_list))
            self.Revise_MakeQuestion_View.ResetPage()
            self.Revise_MakeQuestion_View.show()

            self.hide()
    
    # 接收 調整題目頁面 的資料 之函數 (有幾個參數就接幾個) (bool, bool)
    def GetReviseMakeQuestionViewData(self, is_close, is_return_mainpage = False):
        if is_close == True:
            self.Is_revise_question_view_open = False
            if is_return_mainpage == False: # 不回到首頁
                self.show()
            else: # 回到首頁
                self.return_mainpage_signal.emit(True)

    #####################################################
    #def showEvent(self, event):
    #    self.InputOnTotleNumber()