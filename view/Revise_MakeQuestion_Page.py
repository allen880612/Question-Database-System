from PyQt5.QtWidgets import QMainWindow, QMessageBox
from PyQt5 import QtCore, QtWidgets, QtGui
from view.UI import Revise_MakeQuestion_UI 
from view import ComboboxView as cbview
from model import MyLibrary
from controller import QuestionParser
import random
import os
import docx

class ReviseMakeQuestionPage(QMainWindow):

    revise_make_question_signal = QtCore.pyqtSignal(bool, bool) # set 信號
    TEMPLATE_WORD_PATH = "database/default.docx"

    def __init__(self, _model):
        super(ReviseMakeQuestionPage, self).__init__()
        self.ui = Revise_MakeQuestion_UI.ReviseMakeQuestion_UI()
        self.ui.setupUi(self)
        self.model = _model
        self.is_click_button = False # 是否透過button關閉視窗
        self.is_return_mainpage = False # 是否回到首頁
        # ui層
        self.listWidget_question_level = self.ui.listWidget_make_question_level
        self.listWidget_none_select_question = self.ui.listWidget_none_select_question
        self.listWidget_selected_question = self.ui.listWidget_selected_question
        self.preview_question_content = self.ui.textEdit_preview_question
        self.tbox_search = self.ui.text_search
        
        # 前幾層所選之題目階層 (tuple list: level, question number)
        self.question_level_tupleList = []

        # 未選擇到的題目dict (key: level, value: node list)
        self.question_nonSelect_dict = {}

        # 選擇到的題目dict (key: level, value: node list)
        self.question_select_dict = {}

        # 搜尋條件
        self.search_text = ""

        self.Initialize()

    # 初始化
    def Initialize(self):
        self.ConnectEvent()
        self.ResetPage()

    # 重設頁面
    def ResetPage(self):
        self.is_click_button = False
        self.is_return_mainpage = False

        level_list = [level[0] for level in self.question_level_tupleList]
        number_list = [number[1] for number in self.question_level_tupleList]
        self.UpdateListWidget(self.listWidget_question_level, [MyLibrary.GetQuestionShowText(level) for level in level_list])
        self.listWidget_none_select_question.clear()
        self.listWidget_selected_question.clear()

        self.InitializeDict()
        self.UpdateUI()

    # 註冊事件
    def ConnectEvent(self):
        self.ui.button_return.clicked.connect(self.ClosePage)
        self.ui.listWidget_make_question_level.currentItemChanged.connect(self.SelectQuestionLevel)
        self.ui.listWidget_none_select_question.currentItemChanged.connect(lambda: self.SelectQuestion(0))
        self.ui.listWidget_selected_question.currentItemChanged.connect(lambda: self.SelectQuestion(1))

        self.ui.button_add_question.clicked.connect(self.AddToMakeQuestionList)
        self.ui.button_remove_question.clicked.connect(self.RemoveFromMakeQuestionList)

        self.ui.button_search.clicked.connect(self.InputSearch)
        self.ui.button_continue.clicked.connect(self.MakeQuestion) # Link 出題按鈕

        #self.ui.text_search.textChanged.connect()

    # 刷新UI
    def UpdateUI(self):
        self.ui.button_remove_question.setEnabled(self.listWidget_selected_question.currentRow() > -1)
        self.ui.button_add_question.setEnabled(self.listWidget_none_select_question.currentRow() > -1)

    # 更新 listwidget
    def UpdateListWidget(self, list_widget, data_list):
        list_widget.clear()
        list_widget.addItems(data_list)

    # 更新 未選擇題目的 list widget
    def UpdateNoneSelectQuestionListWidget(self):
        preview_question_count = 15

        level_key = self.GetLevelKey()

        # 搜尋條件
        question_nonSelect_list = self.GetQuestionListBySearch(self.question_nonSelect_dict[level_key], self.search_text)
        
        # 建構 未選擇的題目
        tmp_nonSelectQuestion_list = []
        for question in question_nonSelect_list:
            tmp_nonSelectQuestion_list.append(question.GetQuestion()[:preview_question_count])
        self.UpdateListWidget(self.listWidget_none_select_question, tmp_nonSelectQuestion_list)

    # 更新 已選擇題目的 list widget
    def UpdateSelectQuestionListWidget(self):
        preview_question_count = 15

        level_key = self.GetLevelKey()

        # 建構 已選擇的題目
        tmp_selectQuestion_list = []
        for question in self.question_select_dict[level_key]:
            tmp_selectQuestion_list.append(question.GetQuestion()[:preview_question_count])
        self.UpdateListWidget(self.listWidget_selected_question, tmp_selectQuestion_list)

    # 初始化字典
    def InitializeDict(self):
        self.question_nonSelect_dict.clear()
        self.question_select_dict.clear()

        #level_list = [level[0] for level in self.question_level_tupleList]
        #number_list = [number[1] for number in self.question_level_tupleList]

        for level, number in self.question_level_tupleList:
            keys = MyLibrary.CreateDictKey(level)
            qList = self.model.GetQuestionList(level) # Get Question List

            # 一階初始化
            self.question_nonSelect_dict[keys] = []
            self.question_select_dict[keys] = []

            number_list = [i for i in range(0, len(qList))] # 取得0 ~ len(qList)的數字陣列
            number_list = random.sample(number_list, min(number, len(qList))) # 取得不重複 (number) 個題
            
            # 開建
            for qIndex in range(0, len(qList)):
                if qIndex in number_list: # 被選到
                    self.question_select_dict[keys].append(qList[qIndex])
                else:
                    self.question_nonSelect_dict[keys].append(qList[qIndex])

    # 選擇 題目階層
    def SelectQuestionLevel(self, item):
        preview_question_count = 15
        current_row = self.listWidget_question_level.currentRow()
        if current_row == -1:
            return

        level_key = self.GetLevelKey()

        # 更新 未選擇的題目的 List Widget
        self.UpdateNoneSelectQuestionListWidget()

        # 更新 已選擇的題目的 List Widget
        self.UpdateSelectQuestionListWidget()

        self.UpdateUI()
    
    # 選擇題目
    def SelectQuestion(self, case):
        list_widget = self.sender()
        current_row = list_widget.currentRow()
        if current_row == -1:
            return

        level_key = self.GetLevelKey()

        if case == 0:
            tmp_text = self.question_nonSelect_dict[level_key][current_row].GetQuestion()
            self.preview_question_content.setText(tmp_text)
            self.listWidget_selected_question.setCurrentRow(-1)
        elif case == 1:
            tmp_text = self.question_select_dict[level_key][current_row].GetQuestion()
            self.preview_question_content.setText(tmp_text)
            self.listWidget_none_select_question.setCurrentRow(-1)

        self.UpdateUI()

    # 加入到出題列表
    def AddToMakeQuestionList(self):
        current_row = self.listWidget_none_select_question.currentRow()
        if current_row == -1:
            return

        level_key = self.GetLevelKey()

        question_nonSelect_list = self.question_nonSelect_dict[level_key]
        question_nonSelect_list = self.GetQuestionListBySearch(question_nonSelect_list, self.search_text)

        select_question = question_nonSelect_list[current_row] # 取得題目
        self.question_select_dict[level_key].append(select_question) # 已選擇題目 增加一筆
        self.question_nonSelect_dict[level_key].remove(select_question) # 未選擇題目 刪除這筆
        self.listWidget_none_select_question.setCurrentRow(-1) # 取消Focus
        
        # 更新 未選擇的題目的 List Widget
        self.UpdateNoneSelectQuestionListWidget()
        # 更新 已選擇的題目的 List Widget
        self.UpdateSelectQuestionListWidget()

        self.UpdateUI()
        
    # 從出題列表中移除
    def RemoveFromMakeQuestionList(self):
        current_row = self.listWidget_selected_question.currentRow()
        if current_row == -1:
            return

        level_key = self.GetLevelKey()

        select_question = self.question_select_dict[level_key][current_row] # 取得題目
        self.question_nonSelect_dict[level_key].append(select_question) # 未選擇題目 增加一筆
        self.question_select_dict[level_key].remove(select_question) # 已選擇題目 刪除這筆
        self.listWidget_selected_question.setCurrentRow(-1) # 取消Focus

        # 更新 未選擇的題目的 List Widget
        self.UpdateNoneSelectQuestionListWidget()
        # 更新 已選擇的題目的 List Widget
        self.UpdateSelectQuestionListWidget()

        self.UpdateUI()

    # 取得用於字典查詢時的level -> type = tuple
    def GetLevelKey(self):
        level = self.question_level_tupleList[self.listWidget_question_level.currentRow()][0]
        level_key = MyLibrary.CreateDictKey(level)
        return level_key

    # 設置必要的資料 
    def SetQuestionLevelTupleData(self, tuple_list):
        self.question_level_tupleList = tuple_list

    # 關閉這個視窗
    def ClosePage(self):
        self.is_click_button = True
        is_close = self.close()
        is_return_mainpage = self.is_return_mainpage
        self.revise_make_question_signal.emit(is_close, is_return_mainpage)

    # 關閉視窗事件
    def closeEvent(self, event):
        self.question_level_tupleList = []
        if self.is_click_button == False:
            is_close = True
            is_return_mainpage = self.is_return_mainpage
            self.revise_make_question_signal.emit(is_close, is_return_mainpage)
    
    # 得到Question List
    def GetQuestionList(self):
        qList = []
        for level, questionList in self.question_select_dict.items():
            qList += questionList
        return qList

    # 搜尋框 input
    def InputSearch(self):
        tbox = self.tbox_search
        self.search_text = tbox.text()
        print(self.search_text)

        self.UpdateNoneSelectQuestionListWidget()
        self.UpdateUI()
        
    # 得到 question list 藉由 search
    def GetQuestionListBySearch(self, question_list, search):
        if search == "":
            return question_list
        else:
            new_question_list = []
            for question in question_list:
                #print(question.GetQuestion())
                #print(search)
                #print(question.GetQuestion().find(search))
                if question.GetQuestion().find(search) != -1:
                    new_question_list.append(question)
            return new_question_list

    ###################################################################################
    #出題
    #region 函式區
    def MakeQuestion(self):
        #excel
        if not MyLibrary.IskWordOpen("word/answer.docx") or not MyLibrary.IskWordOpen("word/question.docx"):
            QMessageBox.information(self, "警告", "Word開啟中！\n請關閉Word後再試一次", QMessageBox.Yes)
            return

        qList = self.GetQuestionList() # 要給Word建構的題目列表
        random.shuffle(qList)

        self.BulidWord(qList, "answer", True)  # 建造word 保留答案
        self.BulidWord(qList, "question", False)  # 建造word 刪除答案
        MyLibrary.OpenWord("answer.docx")
        MyLibrary.OpenWord("question.docx")
        print("done")
        self.is_return_mainpage = True
        self.ClosePage()
        # return to index
        #excel

    def BulidWord(self, qList, fileName, haveAnswer):
        word = docx.Document(docx=self.TEMPLATE_WORD_PATH)  # 另一個坑，為了讓封裝後也能抓到default.docx，必須指定
        heading = " - ".join(self.question_level_tupleList[0][0])
        word.add_heading(heading, 0) #新增那個醜醜藍字

        #新增大題 style
        mainPhaseStyle = word.styles.add_style("main_phase", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        mainPhaseStyle.font.size = docx.shared.Pt(18)
        mainPhaseStyle.font.name = "Times New Roman"
        mainPhaseStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #??? 過了兩年我還是看不太懂

        # 新增題目 style
        questionStyle = word.styles.add_style("question", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH) #新增一樣式 (樣式名稱, 樣式類型)
        questionStyle.font.size = docx.shared.Pt(12) #更改此樣式的文字大小
        questionStyle.font.name = "Times New Roman" #設定英文字體
        # where am I? who am I???
        questionStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #設定中文字體
        # 設定凸排, su go i ne, my Python

        questionStyle.paragraph_format.first_line_indent = docx.shared.Pt(-18) # 設定首縮排/凸排 (正值 = 縮排, 負值 = 凸排)
        questionStyle.paragraph_format.left_indent = docx.shared.Pt(18) # ↓注意，重點來了，設定"整個段落"縮排  (正常來說應該不用設定，但是設定凸排的時候，他會順便把整個段落也往左移動，所以要他媽的移回來)

        # 新增選擇題題目 style
        selectQuestionStyle = word.styles.add_style("select_question", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH) #新增一樣式 (樣式名稱, 樣式類型)
        selectQuestionStyle.font.size = docx.shared.Pt(12) #更改此樣式的文字大小
        selectQuestionStyle.font.name = "Times New Roman" #設定英文字體
        selectQuestionStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #設定中文字體
        selectQuestionStyle.paragraph_format.first_line_indent = docx.shared.Pt(-42)
        selectQuestionStyle.paragraph_format.left_indent = docx.shared.Pt(42) 

        #新增題目 - 填充題
        filling_question_parser = QuestionParser.QuestionParser()

        if self.HaveFillingQuestion(qList):
            paragraph = word.add_paragraph("一、填充題", style = "main_phase")
            count = 1
        for i in range(0, len(qList)):
            if qList[i].GetType() == "FillingQuestion":
                #questionIndex = "(" + str(count) + ") " #題號
                questionIndex = str(count) + ". "
                if count < 10:
                    questionIndex += '  '

                count += 1
                if haveAnswer:
                    question = (qList[i].GetAnswer())
                else:
                    question = qList[i].Convert2WordContent(qList[i].GetQuestion())

                #paragraph = word.add_paragraph(questionIndex + question, style = "question")
                paragraph = word.add_paragraph( questionIndex, questionStyle )
                filling_question_parser.Initialize( question, paragraph)
                filling_question_parser.ParseQuestion()


                #paragraph = word.add_paragraph(questionIndex + questionList[i], style = "question") #題號 + 題目 一題作為一個段落
                paragraph.alignment = 0 #設定段落對齊 0 = 靠左, 1 = 置中, 2 = 靠右, 3 = 左右對齊 (WD_PARAGRAPH_ALIGNMENT)
                try:
                    # print("have", qList[i].HaveImage())
                    # print("image", qList[i].GetImage())
                    if qList[i].HaveImage():
                        #run = word.paragraphs[i + 1].add_run()
                        imageParagraph = word.add_paragraph() # 為了讓圖片靠右，直接新增一個段落，方便用alignment
                        imageParagraph.alignment = 2
                        run = imageParagraph.add_run()
                        #run.add_break() #不換段換行
                        for image in qList[i].GetImages():
                            run.add_picture(image.GetWordImage(), height=docx.shared.Cm(2.6))
                except:
                    print("Insert image fail!")
        
        # 新增題目 - 選擇題
        select_question_parser = QuestionParser.QuestionParser()
        if self.HaveSelectQuestion(qList):
            paragraph = word.add_paragraph("二、選擇題", style = "main_phase")
        count = 1
        for question in qList:
            if question.GetType() == "SelectQuestion":
                if haveAnswer:
                    answer_area = "( {0} )".format(" ".join(question.GetAnswer()))
                else:
                    answer_area = "(            )"
                #questionIndex = "(" + str(count) + ") "
                questionIndex = str(count) + ". "
                count += 1
                question_area = answer_area + " " + questionIndex
                
                paragraph = word.add_paragraph(question_area, style = "select_question")
                select_question_parser.Initialize( question.GetQuestion(), paragraph)
                select_question_parser.ParseQuestion()

                paragraph.alignment = 0
                run = paragraph.add_run()
                try:
                    if question.HaveImage():
                        run.add_break()
                        for image in question.GetImages():
                            run.add_picture(image.GetWordImage(), height=docx.shared.Cm(2.6))
                except:
                    print("Insert select image fail!")

                # 新增選項
                option_count = 1
                for options in question.option:
                    run.add_break()
                    option = "(" + self.ConvertOptionNumber2English(option_count) + ") "
                    paragraph.add_run(option)

                    # 處理選項敘述中，可能包含的特殊符號
                    select_question_parser.Initialize( options.GetContent(), paragraph)
                    select_question_parser.ParseQuestion()

                    option_count += 1
                    run = paragraph.add_run('')
                    try:
                        if options.HaveImage():
                            run.add_break()
                            for image in options.GetImages():
                                run.add_picture(image.GetWordImage(), height=docx.shared.Cm(2.6))
                    except Exception as e:
                        print("Insert option image fail!")
                        print(e)

        savePath = "word/" + fileName + ".docx"
        word.save(savePath) #存檔 (存在word資料夾)
    #endregion

    # 檢查有沒有填充題
    def HaveFillingQuestion(self, qList):
        for question in qList:
            if question.GetType() == "FillingQuestion":
                return True
        return False

    # 檢查有沒有選擇題
    def HaveSelectQuestion(self, qList):
        for question in qList:
            if question.GetType() == "SelectQuestion":
                return True
        return False

    # 將選項的數字轉換成英文
    def ConvertOptionNumber2English(self, number):
        converter = [' ', 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        return converter[number % 8]