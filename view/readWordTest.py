import os
import docx

#file = docx.Document("../database/d/chapter.docx")

path = "database/基礎化學/一"
# print(os.listdir(path))
# print(os.getcwd())

dirName = os.path.abspath(os.path.dirname(os.getcwd()))
print(dirName)
aim_dir = os.path.join(dirName, path)
print(aim_dir)
files = os.listdir(aim_dir)
print(files)
os.startfile(aim_dir + "/" + "why.docx")

# os.startfile("../database/基礎化學/一/第一章 物質的組成.docx")
# file = docx.Document("../database/基礎化學/一/第一章 物質的組成.docx")
#
# for f in file.paragraphs:
#     print(f.text)
