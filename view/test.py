from PyQt5 import QtCore, QtGui, QtWidgets
from controller import GET_PATH as gp
from controller import GET_QUESTION as gq
from view import ComboboxView as cbview
from model import MyLibrary
import random
import os
import docx
import numpy as np
#from docx import Document
#from docx.shared import Inches

class Ui_MainWindow(object):
    path = ""
    DATABASE = "database/"
    fManager = gp.FolderManager()
    path2 = ["database/"]
    wordPath = ""

    #excel
    excelPath = "database/盈虧問題v2.xlsx"
    excel = gq.ExcelManager(excelPath)
    #endexcel

    #region 創建下拉式選單
    comboboxView = cbview.ComboboxManager(excel.GetOriginalDataFrame())
    comboboxSelectOption = []
    #endregion

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(860, 480)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(60, 50, 571, 71))
        font = QtGui.QFont()
        font.setFamily("Consolas")
        font.setPointSize(26)
        font.setBold(True)
        font.setWeight(75)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.btn_confirm = QtWidgets.QPushButton(self.centralwidget)
        self.btn_confirm.setGeometry(QtCore.QRect(70, 340, 75, 23))
        self.btn_confirm.setObjectName("btn_confirm")
        self.cBox_dir1 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_dir1.setGeometry(QtCore.QRect(60, 140, 121, 21))
        self.cBox_dir1.setObjectName("cBox_dir1")
        self.cBox_dir2 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_dir2.setGeometry(QtCore.QRect(210, 140, 121, 21))
        self.cBox_dir2.setObjectName("cBox_dir2")
        self.cBox_word = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_word.setGeometry(QtCore.QRect(360, 140, 191, 21))
        self.cBox_word.setObjectName("cBox_word")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(60, 120, 91, 16))
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(210, 120, 81, 16))
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(360, 120, 71, 16))
        self.label_4.setObjectName("label_4")
        self.cBox_level2 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_level2.setGeometry(QtCore.QRect(200, 230, 121, 21))
        self.cBox_level2.setObjectName("cBox_level2")
        self.label_level2 = QtWidgets.QLabel(self.centralwidget)
        self.label_level2.setGeometry(QtCore.QRect(200, 205, 71, 21))
        self.label_level2.setObjectName("label_level2")
        self.cBox_level3 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_level3.setGeometry(QtCore.QRect(340, 230, 121, 21))
        self.cBox_level3.setObjectName("cBox_level3")
        self.label_level3 = QtWidgets.QLabel(self.centralwidget)
        self.label_level3.setGeometry(QtCore.QRect(340, 205, 71, 21))
        self.label_level3.setObjectName("label_level3")
        self.cBox_level4 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_level4.setGeometry(QtCore.QRect(480, 230, 121, 21))
        self.cBox_level4.setObjectName("cBox_level4")
        self.label_level4 = QtWidgets.QLabel(self.centralwidget)
        self.label_level4.setGeometry(QtCore.QRect(480, 205, 71, 21))
        self.label_level4.setObjectName("label_level4")
        self.cBox_level5 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_level5.setGeometry(QtCore.QRect(620, 230, 121, 21))
        self.cBox_level5.setObjectName("cBox_level5")
        self.label_level5 = QtWidgets.QLabel(self.centralwidget)
        self.label_level5.setGeometry(QtCore.QRect(620, 205, 71, 21))
        self.label_level5.setObjectName("label_level5")
        self.label_level1 = QtWidgets.QLabel(self.centralwidget)
        self.label_level1.setGeometry(QtCore.QRect(60, 200, 119, 23))
        self.label_level1.setObjectName("label_level1")
        self.cBox_level1 = QtWidgets.QComboBox(self.centralwidget)
        self.cBox_level1.setGeometry(QtCore.QRect(60, 230, 119, 20))
        self.cBox_level1.setObjectName("cBox_level1")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 860, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.InitGUI()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Question Database"))
        self.label.setText(_translate("MainWindow", "請選擇試題範圍"))
        self.btn_confirm.setText(_translate("MainWindow", "確定"))
        self.label_2.setText(_translate("MainWindow", "資料夾 一層"))
        self.label_3.setText(_translate("MainWindow", "資料夾 二層"))
        self.label_4.setText(_translate("MainWindow", "Word"))
        self.label_level1.setText(_translate("MainWindow", "資料夾 一層"))
        self.label_level2.setText(_translate("MainWindow", "資料夾 二層"))
        self.label_level3.setText(_translate("MainWindow", "資料夾 三層"))
        self.label_level4.setText(_translate("MainWindow", "資料夾 四層"))
        self.label_level5.setText(_translate("MainWindow", "資料夾 五層"))

    def InitGUI(self):
        # 先放著，之後應該會用到
        fullDir = os.walk(self.DATABASE)
        for fr in fullDir:
            print(fr)

        self.btn_confirm.clicked.connect(self.Confirm)
        self.path = self.DATABASE
        self.cBox_dir1.addItems(self.fManager.GetNextLevel(self.DATABASE))
        self.cBox_dir1.activated[str].connect(self.RebuildDir2)
        self.cBox_dir2.activated[str].connect(self.RebuildDir3)
        self.cBox_word.activated[str].connect(self.GetWordPath)

        self.cBox_dir1.setEditable(True)
        self.cBox_dir2.setEditable(True)
        self.cBox_word.setEditable(True)
        self.cBox_dir1.lineEdit().setText("請選擇科目")
        self.cBox_dir2.lineEdit().setText("請先選擇科目")
        self.cBox_word.lineEdit().setText("請選擇科目及分類")

        #region 下拉式選單
        defaultString = self.comboboxView.GetNoSelectString()
        self.cBox_level1.addItems(self.comboboxView.GetDictValue(defaultString))
        self.cBox_level1.activated[str].connect(self.SelectLevel1)
        self.cBox_level2.activated[str].connect(self.SelectLevel2)
        self.cBox_level3.activated[str].connect(self.SelectLevel3)
        self.cBox_level4.activated[str].connect(self.SelectLevel4)
        self.cBox_level5.activated[str].connect(self.SelectLevel5)


        self.cBox_level1.setEditable(True)
        self.cBox_level2.setEditable(True)
        self.cBox_level3.setEditable(True)
        self.cBox_level4.setEditable(True)
        self.cBox_level5.setEditable(True)

        self.cBox_level1.lineEdit().setText("選擇第一層")
        self.cBox_level2.lineEdit().setText("選擇第二層")
        self.cBox_level3.lineEdit().setText("選擇第三層")
        self.cBox_level4.lineEdit().setText("選擇第四層")
        self.cBox_level5.lineEdit().setText("選擇第五層")
        #endregion

    def SelectLevel1(self, text):
        #comboboxSelectOption = 目前選到的層級 (之後做成key)
        self.comboboxSelectOption = [self.cBox_level1.currentText()]
        #先刪除後面的下拉選單的items再重新加入
        self.cBox_level2.clear()
        self.cBox_level3.clear()
        self.cBox_level4.clear()
        self.cBox_level5.clear()
        self.cBox_level2.addItems(self.comboboxView.GetDictValue(text))
        self.cBox_level2.setEditText("選擇第二層")
        self.cBox_level3.setEditText("選擇第三層")
        self.cBox_level4.setEditText("選擇第四層")
        self.cBox_level5.setEditText("選擇第五層")

    def SelectLevel2(self, text):
        self.comboboxSelectOption = [self.cBox_level1.currentText(), self.cBox_level2.currentText()]

        self.cBox_level3.clear()
        self.cBox_level4.clear()
        self.cBox_level5.clear()
        self.cBox_level3.addItems(self.comboboxView.GetDictValue(MyLibrary.CreateDictKey(self.comboboxSelectOption)))
        self.cBox_level3.setEditText("選擇第三層")
        self.cBox_level4.setEditText("選擇第四層")
        self.cBox_level5.setEditText("選擇第五層")

    def SelectLevel3(self, text):
        self.comboboxSelectOption = [self.cBox_level1.currentText(), self.cBox_level2.currentText(), self.cBox_level3.currentText()]

        self.cBox_level4.clear()
        self.cBox_level5.clear()
        self.cBox_level4.addItems(self.comboboxView.GetDictValue(MyLibrary.CreateDictKey(self.comboboxSelectOption)))
        self.cBox_level4.setEditText("選擇第四層")
        self.cBox_level5.setEditText("選擇第五層")

    def SelectLevel4(self, text):
        self.comboboxSelectOption = [self.cBox_level1.currentText(), self.cBox_level2.currentText(), self.cBox_level3.currentText(), self.cBox_level4.currentText()]

        self.cBox_level5.clear()
        self.cBox_level5.addItems(self.comboboxView.GetDictValue(MyLibrary.CreateDictKey(self.comboboxSelectOption)))
        self.cBox_level5.setEditText("選擇第五層")

    def SelectLevel5(self, text):
        self.comboboxSelectOption = [self.cBox_level1.currentText(), self.cBox_level2.currentText(), self.cBox_level3.currentText(), self.cBox_level4.currentText(), self.cBox_level5.currentText()]
        print(self.comboboxSelectOption)

    def ClearCombobox(self):
        self.comboboxSelectOption.clear()
        self.cBox_level1.clear()
        self.cBox_level2.clear()
        self.cBox_level3.clear()
        self.cBox_level4.clear()
        self.cBox_level5.clear()

        self.cBox_level1.addItems(self.comboboxView.GetDictValue(self.comboboxView.GetNoSelectString())) # add default items
        self.cBox_level1.lineEdit().setText("選擇第一層")
        self.cBox_level2.lineEdit().setText("選擇第二層")
        self.cBox_level3.lineEdit().setText("選擇第三層")
        self.cBox_level4.lineEdit().setText("選擇第四層")
        self.cBox_level5.lineEdit().setText("選擇第五層")

    def RebuildDir2(self, text):
        dir = []
        print("len 2 = " + str(len(self.path2)))
        self.cBox_dir2.clear()
        # self.path = self.DATABASE
        # self.path = os.path.join(self.path, text)
        # print("dir2 = " + self.path)
        # dir = self.fManager.GetNextLevel(self.path)
        if len(self.path2) < 2:
            self.path2.append(text)
        else:
            self.path2[1] = text
        pd = self.path2[:2]
        print(pd)
        path = os.path.join(*pd)
        print("dir2 = " + path)
        dir = self.fManager.GetNextLevel(path)

        # 判斷路徑是否存在，且為資料夾
        if dir:
            self.cBox_dir2.addItems(dir)
            self.cBox_dir2.setEditText("請選擇科目")
            self.cBox_word.setEditText("請選擇題庫")

    def RebuildDir3(self, text):
        dir = []
        self.cBox_word.clear()
        # finalPath = os.path.join(self.path, text)
        # print("dir3 = " + finalPath)
        # dir = self.fManager.GetNextLevel(finalPath)
        if len(self.path2) < 3:
            self.path2.append(text)
        else:
            self.path2[2] = text
        pd = self.path2[:3]
        print(pd)
        path = os.path.join(*pd)
        print("dir3 = " + path)
        dir = self.fManager.GetNextLevel(path)
        # 判斷路徑是否存在，且為資料夾
        if dir:
            self.cBox_word.addItems(dir)
            self.cBox_word.setEditText("請選擇題庫")

    def GetWordPath(self, text):
        print("Word name : " + text)
        self.label.setText(text)

        if len(self.path2) < 4:
            self.path2.append(text)
        else:
            self.path2[3] = text
        print(self.path2)
        self.wordPath = os.path.join(*self.path2)
        # os.startfile(self.wordPath)
        print("word path = " + self.wordPath)

    def Confirm(self):
        if os.path.isfile(self.wordPath):
            os.startfile(self.wordPath)
            print("open " + self.wordPath)
        else:
            self.label.setText("輸入不完全，或文件已損毀!")

        #excel
        if self.excel.IsLoad():
            #questionType = ["數學"]
            questionType = self.comboboxSelectOption #搜尋的條件
            if questionType: #存在搜尋條件才做
                questionNumber = 10 #預設隨機選10題
                questionAnswerList = self.excel.GetFilteredQuestion(questionType) #取得過濾後的題目
                questionAnswerList = random.sample(questionAnswerList, min(questionNumber, len(questionAnswerList))) # 將題目不重複隨機選擇 k 題 (0 <= k <= 篩選後的題目數量)
                questionList = self.DeleteAnswer(questionAnswerList) #刪除答案(【*】)
                print(questionType)
                print(questionAnswerList)
                print(questionList)
                self.BulidWord(questionAnswerList, "answer") #建造word
                self.BulidWord(questionList, "question")  # 建造word
                self.ClearCombobox() #清除選擇的combobox
                print("done")
            else:
                self.label.setText("未選擇條件")
        #excel

    def BulidWord(self, questionList, fileName):
        word = docx.Document()
        word.add_heading("Database", 0) #新增那個醜醜藍字

        #新增題目 style
        questionStyle = word.styles.add_style("question", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH) #新增一樣式 (樣式名稱, 樣式類型)
        questionStyle.font.size = docx.shared.Pt(12) #更改此樣式的文字大小
        questionStyle.font.name = "Times New Roman" #設定英文字體
        # where am I? who am I???
        questionStyle._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "細明體") #設定中文字體
        # 設定凸排, su go i ne, my Python
        questionStyle.paragraph_format.first_line_indent = docx.shared.Pt(-18) # 設定首縮排/凸排 (正值 = 縮排, 負值 = 凸排)
        questionStyle.paragraph_format.left_indent = docx.shared.Pt(18) # ↓注意，重點來了，設定"整個段落"縮排  (正常來說應該不用設定，但是設定凸排的時候，他會順便把整個段落也往左移動，所以要他媽的移回來)

        imageList = self.excel.GetFilteredImage(self.comboboxSelectOption)
        #print(imageList)
        #print(os.getcwd())
        #新增題目
        for i in range(0, len(questionList)):
            questionIndex = "(" + str(i + 1) + ") " #題號
            paragraph = word.add_paragraph(questionIndex + questionList[i], style = "question")
            #paragraph = word.add_paragraph(questionIndex + questionList[i], style = "question") #題號 + 題目 一題作為一個段落
            paragraph.alignment = 3 #設定段落對齊 0 = 靠左, 1 = 置中, 2 = 靠右, 3 = 左右對齊 (WD_PARAGRAPH_ALIGNMENT)

            try:
                # if not np.isnan(imageList[i]):
                if imageList[i][0] != "NOIMAGE":
                    print(imageList[i])
                    run = word.paragraphs[i + 1].add_run()
                    run.add_break() #不換段換行
                    for image in imageList[i]:
                        run.add_picture(image)
            except:
                print("Insert image fail!")
                # print(imageList[i])
                # run = word.paragraphs[i + 1].add_run()
                # run.add_picture(imageList[i])

        #word.add_picture("database\\數學\\應用題\\典型應用題\\燕尾題\\1.png")
        #word.add_page_break() #應該是強制換頁
        savePath = "word/" + fileName + ".docx"
        word.save(savePath) #存檔 (存在word資料夾)

        # 土法煉鋼，不知道效率怎樣
    def DeleteAnswer(self, answerList):
        questionList = []
        for str in answerList:
            newQuestion = ""
            addMode = True  # Mode = True > add a char, False > add a space
            for ch in str:
                if addMode == True or ch == '】':
                    newQuestion += ch
                elif addMode == False:
                    newQuestion += ' '

                if ch == '【':
                    addMode = False
                elif ch == '】':
                    addMode = True
            questionList.append(newQuestion)

        return questionList